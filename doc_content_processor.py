import glob
import docx


def get_text_between_headings(doc, heading1, heading2):
    start = -1
    end = -1
    for i in range(len(doc.paragraphs)):
        if doc.paragraphs[i].text == heading1:
            start = i
        elif doc.paragraphs[i].text == heading2:
            end = i
            break
    return '\n'.join([doc.paragraphs[i].text for i in range(start+1, end)])


# 获取所有Word文件的路径
file_paths = glob.glob('./*.docx')

# 创建一个txt文件
with open('./output.txt', 'w', encoding='utf-8') as f:
    # 遍历每个Word文件，将其内容写入txt文件
    for file_path in file_paths:
        doc = docx.Document(file_path)
        # text = '\n\n'.join([paragraph.text for paragraph in doc.paragraphs])
        # f.write(text)
        fruit = get_text_between_headings(doc, "Done", "Introspection")
        # print(fruit)
        if not (fruit.startswith('Figure') or fruit.startswith('[')):
            f.write(fruit)
